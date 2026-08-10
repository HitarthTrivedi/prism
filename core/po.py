"""
Prism — reading a purchase order
────────────────────────────────
A PO arrives as a PDF, sometimes a Word file, occasionally just typed into the
body of the mail. Prism pulls out the handful of fields that matter, puts them
next to the quotation, and points at anything that differs.

That comparison is the whole value. A rate quietly reduced between quotation
and PO, a quantity changed, a delivery date nobody can meet — each is a
two-second check now or a difficult conversation in three weeks. Nothing here
accepts an order: it prepares the comparison and a person presses the button,
because this is the second of the two places in the workflow where money moves.

**Scanned POs.** Half of them are photographs of a printout, and there is no
text in a photograph. This module detects that and says so plainly rather than
returning a confidently empty order — see looks_scanned().
"""
from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, field
from datetime import date, datetime
from decimal import Decimal

from .quoting import Quotation, indian_currency, rupees, to_decimal

# Below this many characters, a PDF page is a picture with a caption rather
# than a document. Chosen from real scans: an OCR-less read of a scanned page
# yields a stray header or nothing at all, while any genuine PO page carries
# several hundred characters of address, terms and line items.
_TEXT_PER_PAGE = 120


class POError(Exception):
    """Couldn't read the order — always with what to do next in the text."""


# ── getting text out of the file ──────────────────────────────────────────────

def pdf_text(path: str) -> tuple[str, int]:
    """(text, page count) from a PDF. Empty text is a normal answer here."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise POError(
            "This build can't read PDF files. Copy the order details into the "
            "email body and Prism will read them from there.") from e
    try:
        reader = PdfReader(path)
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as e:
        raise POError(
            f"That PDF couldn't be opened — it may be damaged or password "
            f"protected. ({e})") from e
    return "\n".join(pages).strip(), len(pages)


def docx_text(path: str) -> str:
    try:
        import docx
    except ImportError as e:
        raise POError("This build can't read Word files. Save the order as a "
                      "PDF and attach that instead.") from e
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p.strip()).strip()


def looks_scanned(text: str, pages: int = 1) -> bool:
    """True when a PDF is a photograph of a document rather than a document."""
    return len((text or "").strip()) < _TEXT_PER_PAGE * max(1, pages)


SCANNED_ADVICE = (
    "This purchase order is a scan — a picture of a printed page — so there is "
    "no text in it for Prism to read. Type the PO number, date, quantity and "
    "rate into the boxes and everything else carries on as normal."
)


def text_from(path: str) -> str:
    """Text from whatever the customer attached. Raises POError with advice."""
    extension = os.path.splitext(path)[1].lower()
    if extension == ".pdf":
        text, pages = pdf_text(path)
        if looks_scanned(text, pages):
            raise POError(SCANNED_ADVICE)
        return text
    if extension in (".docx", ".docm"):
        return docx_text(path)
    if extension in (".txt", ".csv", ".eml"):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    if extension in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".heic"):
        raise POError(SCANNED_ADVICE)
    raise POError(
        f"Prism can't read a {extension or 'file'} purchase order. A PDF or a "
        f"Word file works, or paste the details into the email.")


def find_attachment(message) -> str | None:
    """The attachment most likely to be the order.

    Name first — an attachment actually called "PO 4471.pdf" is not a guess —
    then the only PDF if there is exactly one. Two unnamed PDFs is genuinely
    ambiguous and returns nothing rather than opening the wrong one.
    """
    names = list(getattr(message, "attachment_names", []) or [])
    for name in names:
        if re.search(r"\b(po|p\.?o\.?|purchase[\s_-]?order|work[\s_-]?order|wo)\b",
                     name, re.I):
            return name
    pdfs = [n for n in names if n.lower().endswith(".pdf")]
    return pdfs[0] if len(pdfs) == 1 else None


# ── the order itself ──────────────────────────────────────────────────────────

@dataclass
class POLine:
    description: str = ""
    quantity: Decimal = Decimal(0)
    unit: str = ""
    rate: Decimal = Decimal(0)
    amount: Decimal = Decimal(0)

    def settled(self) -> "POLine":
        """Fill in whichever of quantity × rate = amount was left out.

        POs routinely print only two of the three. Deriving the third makes the
        comparison against the quotation possible instead of half blank — but
        only ever by arithmetic, never by asking a model to work it out.
        """
        if self.amount == 0 and self.quantity and self.rate:
            self.amount = rupees(self.quantity * self.rate)
        elif self.rate == 0 and self.quantity and self.amount:
            self.rate = rupees(self.amount / self.quantity)
        return self


@dataclass
class PurchaseOrder:
    number: str = ""
    date: date | None = None
    buyer: str = ""
    lines: list[POLine] = field(default_factory=list)
    delivery_date: date | None = None
    total: Decimal = Decimal(0)
    terms: str = ""
    reference: str = ""          # our quotation number, when they quote it
    source: str = ""             # the file it was read from

    @property
    def computed_total(self) -> Decimal:
        return rupees(sum((l.amount for l in self.lines), Decimal(0)))

    @property
    def value(self) -> Decimal:
        """What the order is worth — the printed total, else the sum of lines.

        The printed total wins because it may legitimately include freight or
        tax the line items do not, and it is the figure both sides will quote
        at each other later.
        """
        return self.total or self.computed_total

    def missing(self) -> list[str]:
        """Fields a person still has to supply. Drives the form Prism shows."""
        gaps = []
        if not self.number:
            gaps.append("PO number")
        if not self.date:
            gaps.append("PO date")
        if not self.lines:
            gaps.append("what was ordered")
        elif not any(l.quantity for l in self.lines):
            gaps.append("quantity")
        if not self.value:
            gaps.append("order value")
        return gaps


# ── extraction ────────────────────────────────────────────────────────────────

_PROMPT = """Read this purchase order and return ONLY a JSON object. No
explanation, no markdown fences, no text before or after.

Use exactly these keys:
{
  "po_number": "",
  "po_date": "DD-MM-YYYY",
  "buyer": "the company placing the order",
  "reference": "the supplier quotation number they refer to, if any",
  "delivery_date": "DD-MM-YYYY or empty",
  "terms": "payment or delivery terms in one short line, or empty",
  "total": "the order total as digits only, no symbol or commas",
  "lines": [
    {"description": "", "quantity": "", "unit": "", "rate": "", "amount": ""}
  ]
}

Rules:
- Copy values exactly as printed. Do NOT calculate, convert or round anything.
- Numbers as plain digits: 138000.00 not "Rs. 1,38,000/-".
- If a field is not printed on the document, use an empty string. Never guess
  a PO number, a date or a rate.
- Include every line item, in the order they appear.

PURCHASE ORDER:
"""


def _json_from(text: str) -> dict:
    """Pull the JSON object out of a model's reply.

    Models wrap JSON in fences and occasionally add a sentence in front of it,
    so the outermost braces are located rather than the whole reply parsed.
    """
    raw = (text or "").strip()
    raw = re.sub(r"^```[a-z]*\s*|\s*```$", "", raw, flags=re.I | re.M).strip()
    start, end = raw.find("{"), raw.rfind("}")
    if start == -1 or end <= start:
        raise POError("Prism couldn't make sense of that purchase order. "
                      "Type the details in and carry on.")
    try:
        return json.loads(raw[start:end + 1])
    except json.JSONDecodeError as e:
        raise POError("Prism couldn't make sense of that purchase order. "
                      "Type the details in and carry on.") from e


_DATE_FORMATS = ("%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d", "%d-%b-%Y", "%d %b %Y",
                 "%d.%m.%Y", "%d-%m-%y", "%d/%m/%y")


def parse_date(text: str) -> date | None:
    text = (text or "").strip()
    if not text:
        return None
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue
    return None


def from_json(data: dict, source: str = "") -> PurchaseOrder:
    """Build a PurchaseOrder from the extracted fields, defensively.

    Every value is treated as untrusted text: a model can return a number where
    a string was asked for, a null, or a line item that is a bare string. None
    of those should raise — the caller wants a partly filled form to correct,
    not a traceback.
    """
    order = PurchaseOrder(source=source)
    order.number = str(data.get("po_number") or "").strip()
    order.date = parse_date(str(data.get("po_date") or ""))
    order.buyer = str(data.get("buyer") or "").strip()
    order.reference = str(data.get("reference") or "").strip()
    order.delivery_date = parse_date(str(data.get("delivery_date") or ""))
    order.terms = str(data.get("terms") or "").strip()
    order.total = to_decimal(data.get("total"))

    for raw in (data.get("lines") or []):
        if isinstance(raw, str):
            order.lines.append(POLine(description=raw.strip()))
            continue
        if not isinstance(raw, dict):
            continue
        line = POLine(
            description=str(raw.get("description") or "").strip(),
            quantity=to_decimal(raw.get("quantity")),
            unit=str(raw.get("unit") or "").strip(),
            rate=to_decimal(raw.get("rate")),
            amount=to_decimal(raw.get("amount")),
        ).settled()
        if line.description or line.quantity or line.amount:
            order.lines.append(line)
    return order


def extract(text: str, api_key: str, model: str = "", *,
            source: str = "") -> PurchaseOrder:
    """Read a purchase order out of its text.

    The model's only job is to find the fields. It is told, twice, not to
    calculate anything — every derived figure comes from POLine.settled() and
    PurchaseOrder.computed_total, in Decimal, in Python.
    """
    if not (text or "").strip():
        raise POError(SCANNED_ADVICE)
    if not api_key:
        raise POError("Prism needs its Groq API key set up before it can read "
                      "a purchase order. Setup → Groq API key.")
    from .router import groq_chat
    reply = groq_chat(api_key, model or "", _PROMPT + text[:12000],
                      temperature=0.0, timeout=60)
    return from_json(_json_from(reply), source=source)


# ── comparing it to what we quoted ───────────────────────────────────────────

MONEY = "money"
NOTE = "note"


@dataclass
class Difference:
    field: str
    quoted: str
    ordered: str
    kind: str = MONEY

    @property
    def line(self) -> str:
        return f"{self.field}: quoted {self.quoted}, order says {self.ordered}"


def compare(order: PurchaseOrder, quote: Quotation,
            *, tolerance: Decimal = Decimal("1")) -> list[Difference]:
    """What changed between our quotation and their order.

    Money differences first, because those are the ones that must be seen. A
    tolerance of one rupee absorbs the rounding that happens when a buyer's
    system recomputes tax, without hiding a real reduction.

    An empty list does not mean "accept" — it means there is nothing here a
    person needs to be warned about before they do.
    """
    out: list[Difference] = []

    ordered_value = order.value
    quoted_value = quote.total
    if ordered_value and abs(ordered_value - quoted_value) > tolerance:
        out.append(Difference(
            "Order value", f"₹{indian_currency(quoted_value)}",
            f"₹{indian_currency(ordered_value)}", MONEY))

    # Line by line, in the order they were quoted. Matching by position rather
    # than by description because buyers retype descriptions in their own
    # words, and a name-based match would report every line as changed.
    for index, quoted_line in enumerate(quote.lines):
        if index >= len(order.lines):
            out.append(Difference(
                f"Line {index + 1} ({quoted_line.description[:40]})",
                f"{quoted_line.quantity} {quoted_line.unit}", "not on the order",
                MONEY))
            continue
        ordered_line = order.lines[index]
        if ordered_line.quantity and ordered_line.quantity != quoted_line.quantity:
            out.append(Difference(
                f"Line {index + 1} quantity ({quoted_line.description[:30]})",
                f"{quoted_line.quantity}", f"{ordered_line.quantity}", MONEY))
        # A rate is compared by what the change is WORTH, not by the size of
        # the gap. Ninety paise off a unit rate looks like nothing next to a
        # one-rupee tolerance — and on five thousand pieces it is ₹4,500. The
        # tolerance has to apply to the money, so it is multiplied out first.
        if ordered_line.rate:
            gap = abs(ordered_line.rate - quoted_line.rate)
            count = ordered_line.quantity or quoted_line.quantity or Decimal(1)
            if gap * count > tolerance:
                out.append(Difference(
                    f"Line {index + 1} rate ({quoted_line.description[:30]})",
                    f"₹{indian_currency(quoted_line.rate)}",
                    f"₹{indian_currency(ordered_line.rate)}"
                    f" — ₹{indian_currency(gap * count)} on {count}", MONEY))

    if len(order.lines) > len(quote.lines):
        extra = len(order.lines) - len(quote.lines)
        out.append(Difference("Extra lines", f"{len(quote.lines)} quoted",
                              f"{len(order.lines)} ordered — {extra} not quoted",
                              MONEY))

    # Not money, but the one that causes trouble on the shop floor.
    if order.delivery_date:
        promised = quote.terms.delivery
        out.append(Difference("Delivery", promised or "not stated",
                              order.delivery_date.strftime("%d-%m-%Y"), NOTE))
    if order.reference and quote.number and \
            order.reference.strip().upper() != quote.number.strip().upper():
        out.append(Difference("Quotation referred to", quote.number,
                              order.reference, NOTE))

    out.sort(key=lambda d: 0 if d.kind == MONEY else 1)
    return out


def summary(order: PurchaseOrder, differences: list[Difference]) -> str:
    """One paragraph for the confirmation screen, in plain words."""
    head = (f"Purchase order {order.number or '(number not printed)'} "
            f"from {order.buyer or 'the customer'}"
            f"{' dated ' + order.date.strftime('%d-%m-%Y') if order.date else ''}, "
            f"worth ₹{indian_currency(order.value)}.")
    money = [d for d in differences if d.kind == MONEY]
    if not money:
        return head + " It matches the quotation."
    if len(money) == 1:
        return head + " One thing differs from the quotation — " + money[0].line + "."
    lines = "\n".join(f"  · {d.line}" for d in money)
    return head + f" {len(money)} things differ from the quotation:\n{lines}"
